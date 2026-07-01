from __future__ import annotations
from pathlib import Path
from docx import Document as DocxDocument
from app.parsers.base import BaseParser
from app.models import RawDocument, Page, Table

LINES_PER_PAGE = 40


class DocxParser(BaseParser):
    extensions = [".docx"]

    def parse(self, file_path: Path) -> RawDocument:
        doc = DocxDocument(file_path)
        pages: list[Page] = []
        para_index = 0
        tbl_index = 0

        current_lines: list[str] = []
        current_tables: list[Table] = []
        current_page = 1
        total_lines = 0

        body = doc.element.body
        paragraphs = list(doc.paragraphs)
        tables = list(doc.tables)

        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "p":
                if para_index >= len(paragraphs):
                    continue
                para = paragraphs[para_index]
                para_index += 1

                text = para.text.strip()

                if not text:
                    current_lines.append("")
                    total_lines += 1
                    continue

                heading_level = self._heading_level(para)
                if heading_level is not None:
                    if total_lines > 0:
                        pages.append(self._make_page(current_page, current_lines, current_tables))
                        current_page += 1
                        current_lines = []
                        current_tables = []
                        total_lines = 0
                    prefix = "#" * min(heading_level, 3)
                    current_lines.append(f"{prefix} {text}")
                else:
                    current_lines.append(text)

                total_lines += 1

            elif tag == "tbl":
                if tbl_index >= len(tables):
                    continue
                table = self._parse_table(tables[tbl_index], current_page)
                tbl_index += 1
                current_tables.append(table)
                total_lines += len(table.rows) + 2

            if total_lines >= LINES_PER_PAGE and (current_lines or current_tables):
                header_line = current_lines[0] if current_lines else ""
                if not header_line.startswith("#"):
                    pages.append(self._make_page(current_page, current_lines, current_tables))
                    current_page += 1
                    current_lines = []
                    current_tables = []
                    total_lines = 0

        if current_lines or current_tables:
            pages.append(self._make_page(current_page, current_lines, current_tables))

        if not pages:
            pages.append(Page(number=1, text="(Leeres Dokument)"))

        return RawDocument(
            pages=pages,
            metadata={
                "source_file": file_path.name,
                "format": "docx",
                "page_count": len(pages),
            },
        )

    @staticmethod
    def _heading_level(para) -> int | None:
        style_name = (para.style.name if para.style else "").lower()
        if style_name.startswith("heading") or style_name.startswith("überschrift"):
            try:
                return int("".join(c for c in style_name if c.isdigit()) or "1")
            except ValueError:
                return 1
        if "title" in style_name:
            return 1
        if "subtitle" in style_name:
            return 2
        return None

    @staticmethod
    def _make_page(num: int, lines: list[str], tables: list[Table]) -> Page:
        text = "\n".join(lines).strip()
        return Page(number=num, text=text, tables=tables)

    @staticmethod
    def _parse_table(tbl, page_num: int) -> Table:
        headers: list[str] = []
        rows: list[dict] = []
        for i, row in enumerate(tbl.rows):
            cells = [cell.text.strip() for cell in row.cells]
            if i == 0:
                headers = cells
            else:
                row_dict = {}
                for j, cell in enumerate(cells):
                    key = headers[j] if j < len(headers) else f"col_{j}"
                    row_dict[key] = cell
                rows.append(row_dict)
        return Table(headers=headers, rows=rows, page=page_num)
